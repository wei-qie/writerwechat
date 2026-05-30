"""文章生成引擎：合并数据+模板，输出 Markdown + Word"""

import os
import random
import re
from datetime import datetime
from pathlib import Path
from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from .config import TEMPLATE_DIR, OUTPUT_DIR, WEEKDAY_CN
from .fetchers.a_share import format_index_change

_SUMMARIES = [
    "三大指数集体收跌，资金转向防御板块。",
    "指数分化明显，科技承压，消费走强。",
    "市场放量调整，谨慎观望情绪浓厚。",
    "大盘震荡收红，结构性行情延续。",
    "个股跌多涨少，热点轮动加快。",
    "权重股护盘，题材股熄火，分化格局延续。",
    "增量资金入场意愿不足，短期以震荡为主。",
    "赚钱效应一般，关注量能变化。",
]


def _rand_summary(indices):
    if not indices:
        return random.choice(_SUMMARIES)
    avg_pct = sum((i.get("change_pct") or 0) for i in indices) / len(indices)
    if avg_pct > 0.5:
        return "三大指数集体收涨，市场情绪回暖。"
    elif avg_pct > 0:
        return "三大指数小幅收涨，个股涨跌参半。"
    elif avg_pct > -0.5:
        return "三大指数小幅收跌，市场整体偏弱。"
    else:
        return "三大指数集体收跌，资金转向防御板块。"


class ArticleGenerator:
    def __init__(self):
        self.env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))
        self.slot_names = {
            "morning": "盘前预览",
            "midday": "午间收盘",
            "close": "收盘总结",
            "us_close": "美股收盘",
        }

    def generate(self, slot: str, data: dict) -> str:
        template_map = {
            "morning": "morning.md.j2",
            "midday": "midday.md.j2",
            "close": "close.md.j2",
            "us_close": "us_close.md.j2",
        }
        tmpl_name = template_map.get(slot)
        if not tmpl_name:
            raise ValueError(f"未知时段: {slot}")

        now = datetime.now()
        date_str = now.strftime("%Y年%m月%d日")
        weekday = WEEKDAY_CN[now.weekday()]

        open_desc = data.get("open_desc") or format_index_change(data.get("indices", []))

        tmpl = self.env.get_template(tmpl_name)
        article = tmpl.render(
            date_str=date_str,
            weekday=weekday,
            indices=data.get("indices", []),
            stats=data.get("stats", {"up": 0, "down": 0, "flat": 0, "limit_up": 0, "limit_down": 0, "total": 0}),
            sectors_gain=data.get("sectors_gain", []),
            sectors_lose=data.get("sectors_lose", []),
            fund_gain=data.get("fund_gain", []),
            fund_lose=data.get("fund_lose", []),
            us_indices=data.get("us_indices", []),
            asia_indices=data.get("asia_indices", []),
            summary=_rand_summary(data.get("indices", [])),
            open_desc=open_desc,
        )

        # 清理多余空行
        lines = [line.rstrip() for line in article.split("\n")]
        cleaned = []
        prev_empty = False
        for line in lines:
            if line == "":
                if not prev_empty:
                    cleaned.append(line)
                prev_empty = True
            else:
                cleaned.append(line)
                prev_empty = False
        return "\n".join(cleaned).strip()

    def save(self, slot: str, content: str) -> str:
        """保存 .docx，返回文件路径"""
        now = datetime.now()
        date_dir = now.strftime("%Y-%m-%d")
        out_dir = Path(OUTPUT_DIR) / date_dir
        out_dir.mkdir(parents=True, exist_ok=True)

        basename = f"{self.slot_names[slot]}_{now.strftime('%Y%m%d')}"
        docx_path = out_dir / f"{basename}.docx"
        self._to_docx(content, str(docx_path))
        return str(docx_path)

    def _to_docx(self, md_text: str, output_path: str):
        """将 Markdown 转为 Word 文档"""
        doc = Document()
        # 设置默认中文字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "微软雅黑"
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        lines = md_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            stripped = line

            if not stripped:
                i += 1
                continue

            # 标题 # 或 ##
            if stripped.startswith("## "):
                self._add_heading(doc, stripped[3:], 2)
            elif stripped.startswith("# "):
                self._add_heading(doc, stripped[2:], 1)

            # 引用 >
            elif stripped.startswith("> "):
                self._add_paragraph(doc, stripped[2:], italic=True)

            # 表格行
            elif "|" in stripped and stripped.startswith("|"):
                rows = self._collect_table(lines, i)
                if rows:
                    self._add_table(doc, rows)
                    i += len(rows) - 1

            # 列表项
            elif stripped.startswith("- "):
                self._add_paragraph(doc, stripped[2:], bullet=True)

            # 分隔符 ---
            elif stripped == "---":
                self._add_separator(doc)

            # 纯文本
            else:
                self._add_paragraph(doc, stripped)

            i += 1

        doc.save(output_path)

    def _add_heading(self, doc, text, level):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def _add_paragraph(self, doc, text, italic=False, bullet=False):
        if not text:
            return
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
        else:
            p = doc.add_paragraph()
        self._add_formatted_run(p, text)
        if italic:
            for run in p.runs:
                run.italic = True
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(2)

    def _add_formatted_run(self, paragraph, text):
        """解析 **bold** 标记，添加带格式的 run"""
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                run = paragraph.add_run(part)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def _collect_table(self, lines, start):
        """收集连续的表格行"""
        rows = []
        for j in range(start, len(lines)):
            line = lines[j].strip()
            if line and line.startswith("|"):
                # 跳过对齐行 |:---|:---:|
                if re.search(r"^[\|:\-\s]+$", line):
                    continue
                cells = [c.strip() for c in line.split("|")[1:-1]]
                rows.append(cells)
            else:
                break
        return rows

    def _add_table(self, doc, rows):
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Light Shading Accent 1"
        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < len(rows[r_idx]):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = "微软雅黑"
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        doc.add_paragraph()  # 表后空行

    def _add_separator(self, doc):
        p = doc.add_paragraph()
        run = p.add_run("—" * 30)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
